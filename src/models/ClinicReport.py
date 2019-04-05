import time

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.Elements.MessageBoxes import MessageBoxes
from src.models.DatabaseModel import Database
from src.models.SessionWrapper import SessionWrapper
from src.models.SharedFunctions import SharedFunctions


class ClinicReport:
    def __init__(self, start_date, end_date, show_logo, path_to_save):
        document = Document()
        Tstyle = document.styles['Normal']
        font = Tstyle.font
        font.name = 'Arial'
        font.size = Pt(32)
        obj_styles = document.styles

        # obj_charstyle = obj_styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
        obj_charstyle = obj_styles.add_style('rtl_CHARACTER', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.rtl = True

        obj_proghstyle = obj_styles.add_style('rtl_PARAGRAPH', WD_STYLE_TYPE.PARAGRAPH)
        obj_font_s = obj_proghstyle.font
        obj_font_s.rtl = True

        obj_proghstyle_h = obj_styles.add_style('rtl_HEADING', WD_STYLE_TYPE.PARAGRAPH)
        obj_font_h = obj_proghstyle_h.font
        obj_font_h.rtl = True
        obj_font_h.size = Pt(40)

        obj_charstyle = obj_styles.add_style('JafarStyle', WD_STYLE_TYPE.PARAGRAPH)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(40)
        obj_font.name = 'Times New Roman'
        clinic_id = SessionWrapper.clinic_id
        logo = SharedFunctions.get_logo_img(clinic_id)
        all_doctors = Database().select_all_doctors(clinic_id)
        doctors_names = []
        doctors_counts = {}
        for doctor in all_doctors:
            doctors_names.append(doctor['doctor_name'])
            doctors_counts[doctor['doctor_name']] = 0
        visits = Database().get_visits_for_clinic_report(clinic_id, start_date, end_date)
        number_of_visits = len(visits)
        scheduled_visits = 0
        accomplished_visits = 0
        cancelled_visits = 0
        for visit in visits:
            doctor_name = visit['doctor_name']
            doctors_counts[doctor_name] += 1
            if visit['visit_status'] == "مجدولة":
                scheduled_visits += 1
            if visit['visit_status'] == "كاملة":
                accomplished_visits += 1
            if visit['visit_status'] == "ملغية":
                cancelled_visits += 1
        clinic_info_d = Database().get_clinic_info(clinic_id)

        a = document.add_paragraph(clinic_info_d["clinic_name"], style='rtl_HEADING')
        a.alignment = 1

        if logo and show_logo:
            b = document.add_picture(logo, width=Inches(1.25))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = document.add_paragraph(clinic_info_d['specialization'], style='rtl_HEADING')
        c.alignment = 1
        d = document.add_paragraph(clinic_info_d['address'], style='rtl_HEADING')
        d.alignment = 1
        e = document.add_paragraph("رقم الهاتف "+clinic_info_d['phone'], style='rtl_HEADING')
        e.alignment = 1


        title = 'تقرير حالة المؤسسة'
        p = document.add_paragraph(title, style='rtl_PARAGRAPH')
        p.alignment = 1
        p.bold = True
        title2 = ' للفترة من %s الى %s' % (start_date, end_date)
        sp_a = document.add_paragraph(title2, style='rtl_PARAGRAPH')
        sp_a.alignment = 1
        # sp_a.add_run(title2, style='JafarStyle')
        # sp_a.style = 'Normal'
        # sp_a.add_run(title2, style='Normal') number_of_visits

        sp_1 = document.add_paragraph('العدد الكلي للمراجعات :', style='rtl_PARAGRAPH')
        sp_1.add_run(' ' + str(number_of_visits))
        sp_1.alignment = 2

        sp = document.add_paragraph('عدد المراجعات المنجزة:', style='rtl_PARAGRAPH')
        sp.add_run(' '+str(accomplished_visits))
        sp.alignment = 2
        # sp.style.font.size = Pt(20)
        # sp.style.font.name = 'Times New Roman'

        if scheduled_visits:
            s_o = document.add_paragraph('عدد المراجعات المجدولة:', style='rtl_PARAGRAPH')
            s_o.add_run(' '+str(scheduled_visits))
            s_o.alignment = 2
            s_o.style.font.size = Pt(20)
            s_o.style.font.name = 'Times New Roman'

        s_n = document.add_paragraph('عدد المراجعات الملغية:', style='rtl_PARAGRAPH')
        s_n.add_run(' '+str(cancelled_visits))
        s_n.alignment = 2
        s_n.style.font.size = Pt(20)
        s_n.style.font.name = 'Times New Roman'

        sorted_doctors_counts = sorted(doctors_counts.items(), key=lambda kv: kv[1], reverse=True)
        for row in sorted_doctors_counts:
            s_n = document.add_paragraph(row[0], style='rtl_PARAGRAPH')
            s_n.add_run(': ' + str(row[1]))
            s_n.alignment = 2
        document.add_page_break()
        try:
            time_stamp = time.time()
            file_name = 'clinic_report_'+str(time_stamp)+'.docx'
            document.save(path_to_save+'\\'+file_name)
            SharedFunctions.open_word(path_to_save+'\\'+file_name)
        except Exception as e:
            MessageBoxes.warning_message("خطأ", str(e))
            return
