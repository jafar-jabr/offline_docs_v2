import time

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.Elements.MessageBoxes import MessageBoxes
from src.models.DatabaseModel import Database
from src.models.SessionWrapper import SessionWrapper
from src.models.SharedFunctions import SharedFunctions


class PatientReport:
    def __init__(self, patient_id, show_logo, path_to_save):

        document = Document()
        clinic_id = SessionWrapper.clinic_id
        all_visits = Database().select_patient_visits_for_report(patient_id)
        logo = SharedFunctions.get_logo_img(clinic_id)
        clinic_info_d = Database().get_clinic_info(clinic_id)

        details = Database().get_patient_by_id(patient_id)

        obj_styles = document.styles

        # obj_charstyle = obj_styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
        obj_charstyle = obj_styles.add_style('rtl_CHARACTER', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.rtl = True

        obj_proghstyle = obj_styles.add_style('rtl_PARAGRAPH', WD_STYLE_TYPE.PARAGRAPH)
        obj_font_s = obj_proghstyle.font
        obj_font_s.rtl = True

        obj_proghstyle_h = obj_styles.add_style('rtl_HEADING', WD_STYLE_TYPE.PARAGRAPH)
        obj_font_h = obj_proghstyle_h.font
        obj_font_h.rtl = True
        obj_font_h.size = Pt(40)

        a = document.add_heading(clinic_info_d["clinic_name"], 0)
        a.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo and show_logo:
            b = document.add_picture(logo, width=Inches(1.25))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = document.add_heading(clinic_info_d['specialization'], level=1)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        d = document.add_heading(clinic_info_d['address'], level=2)
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        e = document.add_heading("رقم الهاتف "+clinic_info_d['phone'], level=3)
        e.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph("", style="rtl_PARAGRAPH")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(details["patient_name"]).bold = True

        age = SharedFunctions.calculate_age(details["birth_date"])
        document.add_paragraph('    العمر : ' + age, style="rtl_PARAGRAPH").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_paragraph('    الطول (سم): ' + str(details["tall"]), style="rtl_PARAGRAPH").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_paragraph('    الوزن (كغم): ' + str(details["weight"]), style="rtl_PARAGRAPH").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if details["phone"]:
            document.add_paragraph('   الهاتف : ' + details["phone"], style="rtl_PARAGRAPH").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if details["about"]:
            sp = document.add_paragraph(str(details["about"]), style="rtl_PARAGRAPH")
        # sp.alignment = 1
        pp = document.add_paragraph("", style='Intense Quote')
        pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pp.add_run('المراجعات: ').bold = True
        for visit in all_visits:
            v = document.add_paragraph("", style='List Number').alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # v.style.font.size = Pt(16)
            # v.style.font.name = 'Times New Roman'
            a1 = document.add_paragraph(" التاريخ : "+visit["visit_date"], style="rtl_PARAGRAPH")
            a1.add_run("   الوقت : "+SharedFunctions.readableTime(visit["visit_time"]), style="rtl_CHARACTER")
            a1.add_run("   الحالة : "+visit["visit_status"], style="rtl_CHARACTER")
            a1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            a2_a = document.add_paragraph("اسم الطبيب:  ", style="rtl_PARAGRAPH")
            a2_b = document.add_paragraph(str(visit["doctor_name"]))
            a2_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            a2_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
            a3_a = document.add_paragraph(":الاعراض ", style="rtl_PARAGRAPH")
            a3_b = document.add_paragraph(str(visit["symptoms"]), style="rtl_PARAGRAPH")
            a3_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            a3_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
            a4_a = document.add_paragraph(":التشخيص ", style="rtl_PARAGRAPH")
            a4_b = document.add_paragraph(str(visit["diagnosis"]), style="rtl_PARAGRAPH")
            a4_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            a4_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
            a5_a = document.add_paragraph(":الوصفة ", style="rtl_PARAGRAPH")
            a5_b = document.add_paragraph(str(visit["prescription"]), style="rtl_PARAGRAPH")
            a5_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            a5_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
            a6_a = document.add_paragraph(":التوصيات ", style="rtl_PARAGRAPH")
            a6_b = document.add_paragraph(str(visit["recommendations"]), style="rtl_PARAGRAPH")
            a6_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            a6_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph("")
        document.add_page_break()
        try:
            time_stamp = time.time()
            file_name = 'patient_report_'+str(time_stamp)+'.docx'
            document.save(path_to_save+'\\'+file_name)
            SharedFunctions.open_word(path_to_save+'\\'+file_name)
        except Exception as e:
            MessageBoxes.warning_message("خطأ", str(e))
            return
