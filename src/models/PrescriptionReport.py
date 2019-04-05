import time

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.Elements.MessageBoxes import MessageBoxes
from src.models.DatabaseModel import Database
from src.models.SharedFunctions import SharedFunctions


class PrescriptionReport:
    def __init__(self, clinic_id, visit_id, selected, show_logo, path_to_save):
        document = Document()

        obj_styles = document.styles

        # obj_charstyle = obj_styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
        obj_charstyle = obj_styles.add_style('rtl_CHARACTER', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.rtl = True

        obj_proghstyle = obj_styles.add_style('rtl_PARAGRAPH', WD_STYLE_TYPE.PARAGRAPH)
        obj_font_s = obj_proghstyle.font
        obj_font_s.rtl = True

        obj_proghstyle_h = obj_styles.add_style('rtl_HEADING', WD_STYLE_TYPE.PARAGRAPH)
        obj_font_h = obj_proghstyle_h.font
        obj_font_h.rtl = True
        obj_font_h.size = Pt(40)

        logo = SharedFunctions.get_logo_img(clinic_id)

        details = Database().get_patient_and_visit(visit_id)

        clinic_info_d = Database().get_clinic_info(clinic_id)

        a = document.add_paragraph(clinic_info_d["clinic_name"], style='rtl_HEADING')
        a.alignment = 1

        if show_logo and logo:
            b = document.add_picture(logo, width=Inches(1.25))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = document.add_paragraph(clinic_info_d['specialization'], style='rtl_HEADING')
        c.alignment = 1
        d = document.add_paragraph(clinic_info_d['address'], style='rtl_HEADING')
        d.alignment = 1
        e = document.add_paragraph("رقم الهاتف "+clinic_info_d['phone'], style='rtl_HEADING')
        e.alignment = 1

        p = document.add_paragraph("", style="rtl_PARAGRAPH")
        p.alignment = 1
        p.add_run(details['doctor_name']).bold = True

        sp = document.add_paragraph(
            details['doctor_specialization']
        )

        p = document.add_paragraph("", style='rtl_PARAGRAPH')
        p.alignment = 2
        p.add_run(details["patient_name"]).bold = True

        age = SharedFunctions.calculate_age(details["birth_date"])
        l1 = document.add_paragraph('    العمر : ' + age, style='rtl_PARAGRAPH')
        l1.add_run('                        ')
        l1.add_run('الجنس:'+details['gender'])
        l1.alignment = 2
        l2 = document.add_paragraph('    الطول (سم): ' + str(details["tall"]))
        l2.add_run('                         ')
        l2.add_run('    الوزن (كغم): ' + str(details["weight"]))
        l2.alignment = 2

        l3 = document.add_paragraph('', style='rtl_PARAGRAPH')
        l3.alignment = 2
        if details["phone"]:
            l3.add_run('   الهاتف : ' + details["phone"]).alignment = 2
            l3.add_run('                        ')
        l3.add_run('العمل:'+ details['occupation'])
        sp.alignment = 1
        if 'prescription' in selected:
            pp = document.add_paragraph("", style='Intense Quote')
            pp.alignment = 2
            pp.add_run('الوصفة: ').bold = True
            if details['prescription']:
                prescriptions = details['prescription'].split(',')
                for pr in prescriptions:
                    document.add_paragraph(pr, style='List Number').alignment = 1

            else:
                document.add_paragraph(
                    "لا يوجد", style='List Bullet'
                ).alignment = 1

        if 'symptoms' in selected:
            sym = document.add_paragraph("", style='Intense Quote')
            sym.alignment = 2
            sym.add_run(':الاعراض ').bold = True
            if details['symptoms']:
                symptoms = details['symptoms'].split(',')
                for symp in symptoms:
                    document.add_paragraph(
                        symp, style='List Bullet'
                    ).alignment = 1
            else:
                document.add_paragraph(
                    "لا يوجد", style='List Bullet'
                ).alignment = 1

        if 'diagnosis' in selected:
            dia = document.add_paragraph("", style='Intense Quote')
            dia.alignment = 2
            dia.add_run('التشخيص: ').bold = True
            if details['diagnosis']:
                diagnosis = details['diagnosis'].split(',')
                for diag in diagnosis:
                    document.add_paragraph(
                        diag, style='List Bullet'
                    ).alignment = 1
            else:
                document.add_paragraph(
                    "لا يوجد", style='List Bullet'
                ).alignment = 1
        if 'recommendations' in selected:
            rec = document.add_paragraph("", style='Intense Quote')
            rec.alignment = 2
            rec.add_run('التوصيات: ').bold = True
            if details['recommendations']:
                recommendations = details['recommendations'].split(',')
                for recom in recommendations:
                    document.add_paragraph(
                        recom, style='List Bullet'
                    ).alignment = 1
            else:
                document.add_paragraph(
                    "لا يوجد", style='List Bullet'
                ).alignment = 1
        document.add_page_break()
        try:
            time_stamp = time.time()
            file_name = 'prescription_'+str(time_stamp)+'.docx'
            document.save(path_to_save+'\\'+file_name)
            SharedFunctions.open_word(path_to_save+'\\'+file_name)
        except Exception as e:
            MessageBoxes.warning_message("خطأ", str(e))
            return
